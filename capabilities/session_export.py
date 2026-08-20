"""Session export (Phase 3, Milestone 6 --
Sentry_AgenticSOC_Capabilities_Brief_for_Antigravity.md).

Any conversation or investigation downloads as PDF, DOCX, Markdown, or
JSON. `sections` bodies are written verbatim -- including whatever
grounding lines (Source:/Client:) are already in them -- so export never
re-summarizes, paraphrases, or drops a source citation.
"""

from __future__ import annotations

import json
import logging
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

REPO_ROOT = Path(__file__).resolve().parent.parent
REPORTS_DIR = REPO_ROOT / "outputs" / "reports"

SUPPORTED_FORMATS = {"markdown", "json", "pdf", "docx"}


@dataclass
class ExportResult:
    kind: str  # "answered" | "execution_error"
    output_path: Optional[Path] = None
    error: Optional[str] = None


def _safe_filename(title: str) -> str:
    cleaned = "".join(c if c.isalnum() or c in "-_ " else "_" for c in title).strip()
    return (cleaned.replace(" ", "_") or "report")[:80]


def export_session(title: str, sections: list[tuple[str, str]], fmt: str) -> ExportResult:
    if fmt not in SUPPORTED_FORMATS:
        return ExportResult(
            kind="execution_error",
            error=f"unsupported format '{fmt}' -- must be one of {sorted(SUPPORTED_FORMATS)}",
        )
    if not sections:
        return ExportResult(kind="execution_error", error="nothing to export -- no sections given")

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    base_name = f"{_safe_filename(title)}_{timestamp}"

    try:
        if fmt == "markdown":
            path = REPORTS_DIR / f"{base_name}.md"
            body = f"# {title}\n\n" + "\n\n".join(f"## {h}\n\n{b}" for h, b in sections)
            path.write_text(body, encoding="utf-8")
        elif fmt == "json":
            path = REPORTS_DIR / f"{base_name}.json"
            path.write_text(
                json.dumps(
                    {
                        "title": title,
                        "generated_at": timestamp,
                        "sections": [{"heading": h, "body": b} for h, b in sections],
                    },
                    indent=2,
                ),
                encoding="utf-8",
            )
        elif fmt == "pdf":
            path = REPORTS_DIR / f"{base_name}.pdf"
            _write_pdf(title, sections, path)
        else:  # docx
            path = REPORTS_DIR / f"{base_name}.docx"
            _write_docx(title, sections, path)
    except Exception as e:  # noqa: BLE001
        logger.error("Session export failed (format=%s): %s", fmt, e, exc_info=True)
        return ExportResult(kind="execution_error", error=f"export failed: {e}")

    return ExportResult(kind="answered", output_path=path)


def _write_pdf(title: str, sections: list[tuple[str, str]], path: Path) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    doc = SimpleDocTemplate(
        str(path), pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18
    )
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1, 0.2 * inch)]
    for heading, body in sections:
        story.append(Paragraph(heading, styles["Heading2"]))
        for line in body.split("\n"):
            if line.strip():
                escaped = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(escaped, styles["BodyText"]))
        story.append(Spacer(1, 0.15 * inch))
    doc.build(story)


def _write_docx(title: str, sections: list[tuple[str, str]], path: Path) -> None:
    import docx

    document = docx.Document()
    document.add_heading(title, level=0)
    for heading, body in sections:
        document.add_heading(heading, level=2)
        for line in body.split("\n"):
            if line.strip():
                document.add_paragraph(line)
    document.save(str(path))
